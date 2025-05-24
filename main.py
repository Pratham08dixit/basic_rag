from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware
from typing import List
from uuid import uuid4
import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_google_genai.embeddings import GoogleGenerativeAIEmbeddings
from langchain.vectorstores import FAISS
from langchain_google_genai.chat_models import ChatGoogleGenerativeAI
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain

import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("GOOGLE_API_KEY not found in .env file")
genai.configure(api_key=api_key)

app = FastAPI()
app.add_middleware(SessionMiddleware, secret_key="your-secret-session-key")  # Change this in production
templates = Jinja2Templates(directory="templates")

# Constants
INDEX_PATH = "faiss_indices"
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".doc", ".docx"}
MAX_FILES = 20
MAX_PAGES_PER_FILE = 1000
MAX_FILE_SIZE_MB = 10

# In-memory session memory store
session_memories = {}

def get_session_memory(session_id: str):
    if session_id not in session_memories:
        session_memories[session_id] = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
    return session_memories[session_id]

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/", response_class=HTMLResponse)
async def upload_and_answer(
    request: Request,
    pdf_files: List[UploadFile] = File(default=[]),
    question: str = Form(default="")
):
    status_messages = []
    answer = None

    session_files = request.session.get("uploaded_files", [])

    # For each incoming file, if it’s new, add it to session_files
    for upload in pdf_files:
        if upload.filename not in session_files:
            session_files.append(upload.filename)

    # Save it back into the session
    request.session["uploaded_files"] = session_files

    uploaded_filenames = [upload.filename for upload in pdf_files]

    if len(pdf_files) > MAX_FILES:
        return templates.TemplateResponse("index.html", {
            "request": request,
            "answer": None,
            "status_messages": [f"Error: You can upload a maximum of {MAX_FILES} files."],
            "uploaded_files": uploaded_filenames,
        })

    if len(pdf_files) > MAX_FILES:
        return templates.TemplateResponse("index.html", {
            "request": request,
            "answer": None,
            "status_messages": [f"Error: You can upload a maximum of {MAX_FILES} files."],
            "uploaded_files": session_files,   # pass the running list
        })

    embedding = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vectorstore = None
    if os.path.exists(INDEX_PATH):
        try:
            vectorstore = FAISS.load_local(INDEX_PATH, embedding, allow_dangerous_deserialization=True)
            status_messages.append("Loaded existing FAISS index.")
        except Exception as e:
            status_messages.append(f"Error loading FAISS index: {e}")

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    all_new_docs = []

    for upload in pdf_files:
        filename = upload.filename
        ext = os.path.splitext(filename)[-1].lower()

        if ext not in ALLOWED_EXTENSIONS:
            status_messages.append(f"Skipped unsupported file: {filename}")
            continue

        size_mb = len(await upload.read()) / (1024 * 1024)
        await upload.seek(0)
        if size_mb > MAX_FILE_SIZE_MB:
            status_messages.append(f"Skipped {filename}: File size exceeds {MAX_FILE_SIZE_MB} MB.")
            continue

        try:
            text = ""

            if ext == ".pdf":
                data = await upload.read()
                await upload.seek(0)
                doc = fitz.open(stream=data, filetype="pdf")
                if len(doc) > MAX_PAGES_PER_FILE:
                    status_messages.append(f"Skipped {filename}: Exceeds {MAX_PAGES_PER_FILE} page limit.")
                    doc.close()
                    continue
                for page in doc:
                    text += page.get_text()
                doc.close()

            elif ext == ".txt":
                content = await upload.read()
                text = content.decode("utf-8", errors="ignore")
                await upload.seek(0)

            elif ext in [".doc", ".docx"]:
                temp_path = f"/tmp/{filename}"
                with open(temp_path, "wb") as f:
                    f.write(await upload.read())
                docx_doc = DocxDocument(temp_path)
                if len(docx_doc.paragraphs) > MAX_PAGES_PER_FILE * 3:  # rough estimate
                    status_messages.append(f"Skipped {filename}: Estimated page count exceeds {MAX_PAGES_PER_FILE}.")
                    continue
                for para in docx_doc.paragraphs:
                    text += para.text + "\n"

            if not text.strip():
                status_messages.append(f"No text found in {filename}.")
                continue

            chunks = text_splitter.split_text(text)
            docs = [Document(page_content=chunk, metadata={"source": filename}) for chunk in chunks]

            if vectorstore:
                vectorstore.add_documents(docs)
                status_messages.append(f"Appended {len(docs)} chunks from {filename} to FAISS index.")
            else:
                all_new_docs.extend(docs)
                status_messages.append(f"Processed {len(docs)} chunks from {filename}.")

        except Exception as e:
            status_messages.append(f"Failed to process {filename}: {e}")

    if all_new_docs and not vectorstore:
        try:
            vectorstore = FAISS.from_documents(all_new_docs, embedding)
            status_messages.append(f"Created new FAISS index with {len(all_new_docs)} documents.")
        except Exception as e:
            status_messages.append(f"Failed to create FAISS index: {e}")

    if vectorstore:
        try:
            vectorstore.save_local(INDEX_PATH)
            status_messages.append("FAISS index saved to disk.")
        except Exception as e:
            status_messages.append(f"Failed to save FAISS index: {e}")

    # --- Session Memory and QnA ---
    session_id = request.session.get("session_id")
    if not session_id:
        session_id = str(uuid4())
        request.session["session_id"] = session_id

    memory = get_session_memory(session_id)

    question = question.strip()
    if question:
        if vectorstore:
            try:
                docs = vectorstore.similarity_search(question, k=5)
                if docs:
                    status_messages.append(f"Retrieved top {len(docs)} relevant chunks for the question.")
                    context = "\n\n".join(doc.page_content for doc in docs)
                else:
                    status_messages.append("No relevant chunks found for the question.")
                    context = ""
                system_prompt = (
                    "Answer the question as detailed as possible from the provided context. "
                    "If the answer is not in the context, say 'answer is not available in the context'."
                )
                user_input = f"Context: {context}\n\nQuestion: {question}"
                llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash")
                messages = [("system", system_prompt), ("human", user_input)]
                ai_message = await llm.ainvoke(messages)
                answer = ai_message.content
                status_messages.append("Generated answer using Gemini.")
            except Exception as e:
                status_messages.append(f"Error during question answering: {e}")
                answer = f"Error: {e}"
        else:
            answer = "answer is not available in the context"
            status_messages.append("Cannot answer question without any indexed documents.")
    else:
        status_messages.append("No question submitted.")

    return templates.TemplateResponse("index.html", {
        "request": request,
        "answer": answer,
        "status_messages": status_messages,
        "uploaded_files": session_files,
    })
